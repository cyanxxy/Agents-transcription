"""Export Agent for handling transcript format conversions"""

import json
import re
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime, timedelta
import xml.etree.ElementTree as ET
from pathlib import Path

from .base_agent import BaseAgent, Message, MessageType


class ExportAgent(BaseAgent):
    """Agent responsible for exporting transcripts to various formats"""
    
    def __init__(self, name: str = "Exporter"):
        super().__init__(name)
        self._capabilities = [
            "export_txt",
            "export_srt",
            "export_vtt",
            "export_json",
            "export_xml",
            "export_docx",
            "export_pdf",
            "export_csv",
            "validate_format",
            "convert_format"
        ]
        
    def get_capabilities(self) -> List[str]:
        """Return list of capabilities"""
        return self._capabilities
    
    async def process_message(self, message: Message) -> Optional[Message]:
        """Process export-related messages"""
        try:
            action = message.content.get("action")
            
            if action == "export":
                format_type = message.content.get("format", "txt").lower()
                if format_type == "txt":
                    return await self._handle_export_txt(message)
                elif format_type == "srt":
                    return await self._handle_export_srt(message)
                elif format_type == "vtt":
                    return await self._handle_export_vtt(message)
                elif format_type == "json":
                    return await self._handle_export_json(message)
                elif format_type == "xml":
                    return await self._handle_export_xml(message)
                elif format_type == "docx":
                    return await self._handle_export_docx(message)
                elif format_type == "pdf":
                    return await self._handle_export_pdf(message)
                elif format_type == "csv":
                    return await self._handle_export_csv(message)
                else:
                    return message.reply(
                        {"error": f"Unsupported format: {format_type}"},
                        MessageType.ERROR
                    )
            elif action == "validate_format":
                return await self._handle_validate_format(message)
            elif action == "convert_format":
                return await self._handle_convert_format(message)
            elif action == "get_supported_formats":
                return await self._handle_get_supported_formats(message)
            else:
                return message.reply(
                    {"error": f"Unknown action: {action}"},
                    MessageType.ERROR
                )
                
        except Exception as e:
            self.logger.error(f"Error in ExportAgent: {e}")
            return message.reply(
                {"error": str(e), "details": "Export failed"},
                MessageType.ERROR
            )
    
    async def _handle_export_txt(self, message: Message) -> Message:
        """Export transcript as plain text"""
        transcript = message.content.get("transcript", "")
        options = message.content.get("options", {})
        
        if not transcript:
            return message.reply(
                {"error": "No transcript provided"},
                MessageType.ERROR
            )
        
        # Apply formatting options
        formatted_text = transcript
        
        if options.get("remove_timestamps", False):
            # Remove timestamp patterns
            formatted_text = re.sub(r'\[\d{2}:\d{2}:\d{2}\.\d{2}\]\s*', '', formatted_text)
        
        if options.get("remove_speaker_labels", False):
            # Remove speaker labels (pattern: "Speaker: ")
            formatted_text = re.sub(r'^(\w+):\s*', '', formatted_text, flags=re.MULTILINE)
        
        if options.get("wrap_lines", False):
            # Wrap long lines
            max_width = options.get("line_width", 80)
            formatted_text = self._wrap_text(formatted_text, max_width)
        
        if options.get("add_header", False):
            # Add header with metadata
            header = self._generate_header(options.get("metadata", {}))
            formatted_text = header + "\n\n" + formatted_text
        
        return message.reply({
            "success": True,
            "content": formatted_text,
            "format": "txt",
            "encoding": "utf-8",
            "size_bytes": len(formatted_text.encode('utf-8'))
        })
    
    async def _handle_export_srt(self, message: Message) -> Message:
        """Export transcript as SRT subtitle file"""
        transcript = message.content.get("transcript", "")
        options = message.content.get("options", {})
        
        if not transcript:
            return message.reply(
                {"error": "No transcript provided"},
                MessageType.ERROR
            )
        
        # Parse transcript and extract segments with timestamps
        segments = self._parse_timestamped_transcript(transcript)
        
        if not segments:
            return message.reply(
                {"error": "No timestamped segments found in transcript"},
                MessageType.ERROR
            )
        
        # Generate SRT content
        srt_lines = []
        for i, segment in enumerate(segments, 1):
            # Format: subtitle number
            srt_lines.append(str(i))
            
            # Format: start_time --> end_time
            start_time = self._format_srt_time(segment["start"])
            end_time = self._format_srt_time(segment["end"])
            srt_lines.append(f"{start_time} --> {end_time}")
            
            # Format: subtitle text
            text = segment["text"]
            
            # Apply options
            if options.get("max_chars_per_line", 0) > 0:
                text = self._break_lines(text, options["max_chars_per_line"])
            
            if options.get("max_lines_per_subtitle", 0) > 0:
                lines = text.split('\n')
                if len(lines) > options["max_lines_per_subtitle"]:
                    # Split into multiple subtitles
                    # This is simplified - production would handle timing better
                    text = '\n'.join(lines[:options["max_lines_per_subtitle"]])
            
            srt_lines.append(text)
            srt_lines.append("")  # Empty line between subtitles
        
        srt_content = '\n'.join(srt_lines)
        
        return message.reply({
            "success": True,
            "content": srt_content,
            "format": "srt",
            "encoding": "utf-8",
            "subtitle_count": len(segments),
            "duration": segments[-1]["end"] if segments else 0
        })
    
    async def _handle_export_vtt(self, message: Message) -> Message:
        """Export transcript as WebVTT file"""
        transcript = message.content.get("transcript", "")
        options = message.content.get("options", {})
        
        if not transcript:
            return message.reply(
                {"error": "No transcript provided"},
                MessageType.ERROR
            )
        
        # Parse transcript
        segments = self._parse_timestamped_transcript(transcript)
        
        if not segments:
            return message.reply(
                {"error": "No timestamped segments found in transcript"},
                MessageType.ERROR
            )
        
        # Generate VTT content
        vtt_lines = ["WEBVTT", ""]
        
        # Add metadata if provided
        if options.get("metadata"):
            for key, value in options["metadata"].items():
                vtt_lines.append(f"NOTE {key}: {value}")
            vtt_lines.append("")
        
        for i, segment in enumerate(segments):
            # Optional cue identifier
            if options.get("include_cue_ids", False):
                vtt_lines.append(f"cue-{i+1}")
            
            # Timestamp line
            start_time = self._format_vtt_time(segment["start"])
            end_time = self._format_vtt_time(segment["end"])
            vtt_lines.append(f"{start_time} --> {end_time}")
            
            # Optional cue settings
            if options.get("position"):
                vtt_lines[-1] += f" position:{options['position']}%"
            
            # Text
            text = segment["text"]
            if options.get("include_speaker") and segment.get("speaker"):
                text = f"<v {segment['speaker']}>{text}"
            
            vtt_lines.append(text)
            vtt_lines.append("")
        
        vtt_content = '\n'.join(vtt_lines)
        
        return message.reply({
            "success": True,
            "content": vtt_content,
            "format": "vtt",
            "encoding": "utf-8",
            "cue_count": len(segments)
        })
    
    async def _handle_export_json(self, message: Message) -> Message:
        """Export transcript as structured JSON"""
        transcript = message.content.get("transcript", "")
        options = message.content.get("options", {})
        
        if not transcript:
            return message.reply(
                {"error": "No transcript provided"},
                MessageType.ERROR
            )
        
        # Parse transcript into structured data
        segments = self._parse_timestamped_transcript(transcript)
        
        # Build JSON structure
        json_data = {
            "version": "1.0",
            "created_at": datetime.now().isoformat(),
            "metadata": options.get("metadata", {}),
            "statistics": {
                "total_segments": len(segments),
                "total_duration": segments[-1]["end"] if segments else 0,
                "word_count": sum(len(s["text"].split()) for s in segments),
                "speaker_count": len(set(s.get("speaker", "Unknown") for s in segments))
            }
        }
        
        # Format segments based on options
        if options.get("hierarchical", False):
            # Group by speaker
            by_speaker = {}
            for segment in segments:
                speaker = segment.get("speaker", "Unknown")
                if speaker not in by_speaker:
                    by_speaker[speaker] = []
                by_speaker[speaker].append(segment)
            json_data["content"] = by_speaker
        else:
            json_data["segments"] = segments
        
        # Add word-level timing if requested
        if options.get("word_timing", False):
            for segment in segments:
                segment["words"] = self._estimate_word_timing(
                    segment["text"], 
                    segment["start"], 
                    segment["end"]
                )
        
        # Convert to JSON string
        json_content = json.dumps(
            json_data, 
            indent=2 if options.get("pretty_print", True) else None,
            ensure_ascii=False
        )
        
        return message.reply({
            "success": True,
            "content": json_content,
            "format": "json",
            "encoding": "utf-8",
            "data": json_data  # Also return parsed data
        })
    
    async def _handle_export_xml(self, message: Message) -> Message:
        """Export transcript as XML"""
        transcript = message.content.get("transcript", "")
        options = message.content.get("options", {})
        
        if not transcript:
            return message.reply(
                {"error": "No transcript provided"},
                MessageType.ERROR
            )
        
        # Parse transcript
        segments = self._parse_timestamped_transcript(transcript)
        
        # Create XML structure
        root = ET.Element("transcript")
        root.set("version", "1.0")
        root.set("created", datetime.now().isoformat())
        
        # Add metadata
        metadata = ET.SubElement(root, "metadata")
        for key, value in options.get("metadata", {}).items():
            elem = ET.SubElement(metadata, key)
            elem.text = str(value)
        
        # Add segments
        segments_elem = ET.SubElement(root, "segments")
        for i, segment in enumerate(segments):
            seg_elem = ET.SubElement(segments_elem, "segment")
            seg_elem.set("id", str(i + 1))
            seg_elem.set("start", str(segment["start"]))
            seg_elem.set("end", str(segment["end"]))
            
            if segment.get("speaker"):
                seg_elem.set("speaker", segment["speaker"])
            
            text_elem = ET.SubElement(seg_elem, "text")
            text_elem.text = segment["text"]
        
        # Convert to string
        xml_content = ET.tostring(
            root, 
            encoding='unicode',
            method='xml'
        )
        
        # Pretty print if requested
        if options.get("pretty_print", True):
            import xml.dom.minidom
            dom = xml.dom.minidom.parseString(xml_content)
            xml_content = dom.toprettyxml(indent="  ")
        
        return message.reply({
            "success": True,
            "content": xml_content,
            "format": "xml",
            "encoding": "utf-8"
        })
    
    async def _handle_export_docx(self, message: Message) -> Message:
        """Export transcript as DOCX (Word document)"""
        transcript = message.content.get("transcript", "")
        options = message.content.get("options", {})
        
        if not transcript:
            return message.reply(
                {"error": "No transcript provided"},
                MessageType.ERROR
            )
        
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return message.reply(
                {"error": "python-docx library not installed"},
                MessageType.ERROR
            )
        
        # Create document
        doc = Document()
        
        # Add title if provided
        if options.get("title"):
            title = doc.add_heading(options["title"], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        if options.get("metadata"):
            doc.add_heading("Metadata", 1)
            for key, value in options["metadata"].items():
                doc.add_paragraph(f"{key}: {value}")
            doc.add_paragraph()  # Empty line
        
        # Add transcript content
        doc.add_heading("Transcript", 1)
        
        # Parse and add segments
        segments = self._parse_timestamped_transcript(transcript)
        
        for segment in segments:
            # Add timestamp if requested
            if options.get("include_timestamps", True):
                p = doc.add_paragraph()
                timestamp_run = p.add_run(f"[{self._seconds_to_time(segment['start'])}] ")
                timestamp_run.font.color.rgb = RGBColor(128, 128, 128)
                timestamp_run.font.size = Pt(10)
            
            # Add speaker if available
            if segment.get("speaker"):
                speaker_run = p.add_run(f"{segment['speaker']}: ")
                speaker_run.bold = True
            
            # Add text
            text_run = p.add_run(segment["text"])
            text_run.font.size = Pt(11)
        
        # Note: In a real implementation, we would save to bytes
        # For now, return a message indicating success
        return message.reply({
            "success": True,
            "format": "docx",
            "message": "DOCX export prepared (actual file generation would require file system access)",
            "segment_count": len(segments)
        })
    
    async def _handle_export_pdf(self, message: Message) -> Message:
        """Export transcript as PDF"""
        transcript = message.content.get("transcript", "")
        options = message.content.get("options", {})
        
        if not transcript:
            return message.reply(
                {"error": "No transcript provided"},
                MessageType.ERROR
            )
        
        # Note: PDF generation would require a library like reportlab
        # For now, return a placeholder response
        
        return message.reply({
            "success": True,
            "format": "pdf",
            "message": "PDF export would require additional libraries (reportlab)",
            "suggested_approach": "Convert to HTML first, then use wkhtmltopdf or similar"
        })
    
    async def _handle_export_csv(self, message: Message) -> Message:
        """Export transcript as CSV"""
        transcript = message.content.get("transcript", "")
        options = message.content.get("options", {})
        
        if not transcript:
            return message.reply(
                {"error": "No transcript provided"},
                MessageType.ERROR
            )
        
        import csv
        import io
        
        # Parse transcript
        segments = self._parse_timestamped_transcript(transcript)
        
        # Create CSV in memory
        output = io.StringIO()
        
        # Define columns
        columns = ["index", "start_time", "end_time", "duration", "speaker", "text"]
        if options.get("include_word_count", False):
            columns.append("word_count")
        
        writer = csv.DictWriter(output, fieldnames=columns)
        
        # Write header
        writer.writeheader()
        
        # Write rows
        for i, segment in enumerate(segments):
            row = {
                "index": i + 1,
                "start_time": self._seconds_to_time(segment["start"]),
                "end_time": self._seconds_to_time(segment["end"]),
                "duration": round(segment["end"] - segment["start"], 2),
                "speaker": segment.get("speaker", ""),
                "text": segment["text"]
            }
            
            if options.get("include_word_count", False):
                row["word_count"] = len(segment["text"].split())
            
            writer.writerow(row)
        
        csv_content = output.getvalue()
        
        return message.reply({
            "success": True,
            "content": csv_content,
            "format": "csv",
            "encoding": "utf-8",
            "row_count": len(segments) + 1  # +1 for header
        })
    
    async def _handle_validate_format(self, message: Message) -> Message:
        """Validate a specific format"""
        content = message.content.get("content", "")
        format_type = message.content.get("format", "").lower()
        
        if not content or not format_type:
            return message.reply(
                {"error": "Content and format are required"},
                MessageType.ERROR
            )
        
        validation_result = {
            "valid": False,
            "errors": [],
            "warnings": []
        }
        
        try:
            if format_type == "srt":
                validation_result = self._validate_srt(content)
            elif format_type == "vtt":
                validation_result = self._validate_vtt(content)
            elif format_type == "json":
                validation_result = self._validate_json(content)
            elif format_type == "xml":
                validation_result = self._validate_xml(content)
            else:
                validation_result["errors"].append(f"Unknown format: {format_type}")
        except Exception as e:
            validation_result["errors"].append(str(e))
        
        validation_result["valid"] = len(validation_result["errors"]) == 0
        
        return message.reply(validation_result)
    
    async def _handle_convert_format(self, message: Message) -> Message:
        """Convert between formats"""
        content = message.content.get("content", "")
        from_format = message.content.get("from_format", "").lower()
        to_format = message.content.get("to_format", "").lower()
        
        if not content or not from_format or not to_format:
            return message.reply(
                {"error": "Content, from_format, and to_format are required"},
                MessageType.ERROR
            )
        
        try:
            # First, parse the input format to a common structure
            if from_format == "srt":
                segments = self._parse_srt(content)
            elif from_format == "vtt":
                segments = self._parse_vtt(content)
            elif from_format == "json":
                segments = self._parse_json_transcript(content)
            else:
                return message.reply(
                    {"error": f"Unsupported source format: {from_format}"},
                    MessageType.ERROR
                )
            
            # Create a temporary transcript from segments
            transcript = self._segments_to_transcript(segments)
            
            # Now export to target format
            export_message = Message(
                content={
                    "action": "export",
                    "format": to_format,
                    "transcript": transcript,
                    "options": message.content.get("options", {})
                }
            )
            
            return await self.process_message(export_message)
            
        except Exception as e:
            return message.reply(
                {"error": f"Conversion failed: {str(e)}"},
                MessageType.ERROR
            )
    
    async def _handle_get_supported_formats(self, message: Message) -> Message:
        """Get list of supported export formats"""
        formats = {
            "txt": {
                "name": "Plain Text",
                "extension": ".txt",
                "description": "Simple text format without formatting",
                "options": ["remove_timestamps", "remove_speaker_labels", "wrap_lines", "add_header"]
            },
            "srt": {
                "name": "SubRip Subtitles",
                "extension": ".srt",
                "description": "Standard subtitle format with timing",
                "options": ["max_chars_per_line", "max_lines_per_subtitle"]
            },
            "vtt": {
                "name": "WebVTT",
                "extension": ".vtt",
                "description": "Web Video Text Tracks format",
                "options": ["include_cue_ids", "position", "include_speaker"]
            },
            "json": {
                "name": "JSON",
                "extension": ".json",
                "description": "Structured data format",
                "options": ["hierarchical", "word_timing", "pretty_print"]
            },
            "xml": {
                "name": "XML",
                "extension": ".xml",
                "description": "Extensible Markup Language format",
                "options": ["pretty_print"]
            },
            "csv": {
                "name": "CSV",
                "extension": ".csv",
                "description": "Comma-separated values for spreadsheets",
                "options": ["include_word_count"]
            },
            "docx": {
                "name": "Microsoft Word",
                "extension": ".docx",
                "description": "Word document format",
                "options": ["title", "include_timestamps"],
                "requires": "python-docx"
            },
            "pdf": {
                "name": "PDF",
                "extension": ".pdf",
                "description": "Portable Document Format",
                "options": ["title", "include_timestamps"],
                "requires": "reportlab",
                "status": "not_implemented"
            }
        }
        
        return message.reply({
            "formats": formats,
            "count": len(formats)
        })
    
    def _parse_timestamped_transcript(self, transcript: str) -> List[Dict[str, Any]]:
        """Parse transcript with timestamps into segments"""
        segments = []
        
        # Pattern for timestamps like [00:01:23.45]
        timestamp_pattern = r'\[(\d{2}):(\d{2}):(\d{2}\.\d{2})\]'
        
        # Split by timestamps
        parts = re.split(timestamp_pattern, transcript)
        
        # Process parts (pattern creates groups: text, hours, minutes, seconds, text, ...)
        i = 0
        while i < len(parts) - 3:
            if i == 0 and parts[i].strip():
                # Text before first timestamp
                segments.append({
                    "start": 0,
                    "end": 0,
                    "text": parts[i].strip(),
                    "speaker": self._extract_speaker(parts[i])
                })
                i += 1
            else:
                # Timestamp parts
                hours = int(parts[i + 1])
                minutes = int(parts[i + 2])
                seconds = float(parts[i + 3])
                
                start_time = hours * 3600 + minutes * 60 + seconds
                
                # Get text after timestamp
                text_start = i + 4
                text = parts[text_start].strip() if text_start < len(parts) else ""
                
                # Find end time (next timestamp or estimate)
                end_time = start_time + 5  # Default 5 seconds
                if text_start + 4 < len(parts):
                    next_hours = int(parts[text_start + 1])
                    next_minutes = int(parts[text_start + 2])
                    next_seconds = float(parts[text_start + 3])
                    end_time = next_hours * 3600 + next_minutes * 60 + next_seconds
                
                if text:
                    segments.append({
                        "start": start_time,
                        "end": end_time,
                        "text": text,
                        "speaker": self._extract_speaker(text)
                    })
                
                i += 4
        
        return segments
    
    def _extract_speaker(self, text: str) -> Optional[str]:
        """Extract speaker name from text if present"""
        # Pattern: "Speaker: text" at the beginning
        match = re.match(r'^(\w+):\s*(.+)', text.strip())
        if match:
            return match.group(1)
        return None
    
    def _format_srt_time(self, seconds: float) -> str:
        """Format time for SRT (00:01:23,456)"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds % 1) * 1000)
        
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"
    
    def _format_vtt_time(self, seconds: float) -> str:
        """Format time for WebVTT (00:01:23.456)"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds % 1) * 1000)
        
        return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millis:03d}"
    
    def _seconds_to_time(self, seconds: float) -> str:
        """Convert seconds to readable time format"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = seconds % 60
        
        return f"{hours:02d}:{minutes:02d}:{secs:05.2f}"
    
    def _wrap_text(self, text: str, width: int) -> str:
        """Wrap text to specified width"""
        import textwrap
        
        lines = text.split('\n')
        wrapped_lines = []
        
        for line in lines:
            if len(line) <= width:
                wrapped_lines.append(line)
            else:
                wrapped = textwrap.wrap(line, width=width)
                wrapped_lines.extend(wrapped)
        
        return '\n'.join(wrapped_lines)
    
    def _break_lines(self, text: str, max_chars: int) -> str:
        """Break text into lines with max characters"""
        words = text.split()
        lines = []
        current_line = []
        current_length = 0
        
        for word in words:
            word_length = len(word)
            if current_length + word_length + 1 <= max_chars:
                current_line.append(word)
                current_length += word_length + 1
            else:
                if current_line:
                    lines.append(' '.join(current_line))
                current_line = [word]
                current_length = word_length
        
        if current_line:
            lines.append(' '.join(current_line))
        
        return '\n'.join(lines)
    
    def _generate_header(self, metadata: Dict[str, Any]) -> str:
        """Generate header for text export"""
        header_lines = [
            "=" * 50,
            "TRANSCRIPT",
            "=" * 50,
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        ]
        
        for key, value in metadata.items():
            header_lines.append(f"{key}: {value}")
        
        header_lines.append("=" * 50)
        
        return '\n'.join(header_lines)
    
    def _estimate_word_timing(self, text: str, start: float, end: float) -> List[Dict[str, Any]]:
        """Estimate word-level timing (simplified)"""
        words = text.split()
        if not words:
            return []
        
        duration = end - start
        word_duration = duration / len(words)
        
        word_timings = []
        current_time = start
        
        for word in words:
            word_timings.append({
                "word": word,
                "start": round(current_time, 3),
                "end": round(current_time + word_duration, 3)
            })
            current_time += word_duration
        
        return word_timings
    
    def _validate_srt(self, content: str) -> Dict[str, Any]:
        """Validate SRT format"""
        result = {"valid": True, "errors": [], "warnings": []}
        
        lines = content.strip().split('\n')
        i = 0
        subtitle_count = 0
        
        while i < len(lines):
            # Skip empty lines
            if not lines[i].strip():
                i += 1
                continue
            
            # Check subtitle number
            if not lines[i].strip().isdigit():
                result["errors"].append(f"Line {i+1}: Expected subtitle number, got '{lines[i]}'")
                break
            
            subtitle_count += 1
            i += 1
            
            # Check timestamp line
            if i >= len(lines):
                result["errors"].append("Unexpected end of file after subtitle number")
                break
            
            timestamp_pattern = r'^\d{2}:\d{2}:\d{2},\d{3} --> \d{2}:\d{2}:\d{2},\d{3}$'
            if not re.match(timestamp_pattern, lines[i].strip()):
                result["errors"].append(f"Line {i+1}: Invalid timestamp format")
                break
            
            i += 1
            
            # Read subtitle text (can be multiple lines)
            text_lines = []
            while i < len(lines) and lines[i].strip():
                text_lines.append(lines[i])
                i += 1
            
            if not text_lines:
                result["warnings"].append(f"Subtitle {subtitle_count} has no text")
        
        if subtitle_count == 0:
            result["errors"].append("No valid subtitles found")
        
        result["valid"] = len(result["errors"]) == 0
        return result
    
    def _validate_vtt(self, content: str) -> Dict[str, Any]:
        """Validate WebVTT format"""
        result = {"valid": True, "errors": [], "warnings": []}
        
        lines = content.strip().split('\n')
        
        # Check header
        if not lines or not lines[0].startswith('WEBVTT'):
            result["errors"].append("File must start with 'WEBVTT'")
            return result
        
        # Basic validation of cues
        i = 1
        while i < len(lines):
            if '-->' in lines[i]:
                # Timestamp line
                timestamp_pattern = r'^\d{2}:\d{2}:\d{2}\.\d{3} --> \d{2}:\d{2}:\d{2}\.\d{3}'
                if not re.match(timestamp_pattern, lines[i].strip()):
                    result["errors"].append(f"Line {i+1}: Invalid timestamp format")
            i += 1
        
        result["valid"] = len(result["errors"]) == 0
        return result
    
    def _validate_json(self, content: str) -> Dict[str, Any]:
        """Validate JSON format"""
        result = {"valid": True, "errors": [], "warnings": []}
        
        try:
            data = json.loads(content)
            
            # Check expected structure
            if "segments" in data:
                for i, segment in enumerate(data["segments"]):
                    if "text" not in segment:
                        result["warnings"].append(f"Segment {i} missing 'text' field")
                    if "start" not in segment:
                        result["warnings"].append(f"Segment {i} missing 'start' field")
            elif "content" not in data:
                result["warnings"].append("No 'segments' or 'content' field found")
                
        except json.JSONDecodeError as e:
            result["errors"].append(f"Invalid JSON: {str(e)}")
        
        result["valid"] = len(result["errors"]) == 0
        return result
    
    def _validate_xml(self, content: str) -> Dict[str, Any]:
        """Validate XML format"""
        result = {"valid": True, "errors": [], "warnings": []}
        
        try:
            root = ET.fromstring(content)
            
            # Check for expected elements
            if root.tag != "transcript":
                result["warnings"].append(f"Root element is '{root.tag}', expected 'transcript'")
            
            segments = root.find("segments")
            if segments is None:
                result["warnings"].append("No 'segments' element found")
            
        except ET.ParseError as e:
            result["errors"].append(f"Invalid XML: {str(e)}")
        
        result["valid"] = len(result["errors"]) == 0
        return result
    
    def _parse_srt(self, content: str) -> List[Dict[str, Any]]:
        """Parse SRT content into segments"""
        segments = []
        lines = content.strip().split('\n')
        i = 0
        
        while i < len(lines):
            # Skip empty lines
            if not lines[i].strip():
                i += 1
                continue
            
            # Skip subtitle number
            i += 1
            
            # Parse timestamp
            if i < len(lines) and '-->' in lines[i]:
                times = lines[i].split('-->')
                start = self._parse_srt_time(times[0].strip())
                end = self._parse_srt_time(times[1].strip())
                i += 1
                
                # Collect text lines
                text_lines = []
                while i < len(lines) and lines[i].strip():
                    text_lines.append(lines[i])
                    i += 1
                
                segments.append({
                    "start": start,
                    "end": end,
                    "text": '\n'.join(text_lines)
                })
        
        return segments
    
    def _parse_srt_time(self, time_str: str) -> float:
        """Parse SRT time format to seconds"""
        parts = re.match(r'(\d{2}):(\d{2}):(\d{2}),(\d{3})', time_str)
        if parts:
            hours = int(parts.group(1))
            minutes = int(parts.group(2))
            seconds = int(parts.group(3))
            millis = int(parts.group(4))
            
            return hours * 3600 + minutes * 60 + seconds + millis / 1000
        
        return 0
    
    def _parse_vtt(self, content: str) -> List[Dict[str, Any]]:
        """Parse WebVTT content into segments"""
        segments = []
        lines = content.strip().split('\n')
        i = 1  # Skip WEBVTT header
        
        while i < len(lines):
            if '-->' in lines[i]:
                # Parse timestamp
                times = lines[i].split('-->')
                start = self._parse_vtt_time(times[0].strip())
                end = self._parse_vtt_time(times[1].strip())
                i += 1
                
                # Collect text lines
                text_lines = []
                while i < len(lines) and lines[i].strip() and '-->' not in lines[i]:
                    text_lines.append(lines[i])
                    i += 1
                
                segments.append({
                    "start": start,
                    "end": end,
                    "text": '\n'.join(text_lines)
                })
            else:
                i += 1
        
        return segments
    
    def _parse_vtt_time(self, time_str: str) -> float:
        """Parse VTT time format to seconds"""
        parts = re.match(r'(\d{2}):(\d{2}):(\d{2})\.(\d{3})', time_str)
        if parts:
            hours = int(parts.group(1))
            minutes = int(parts.group(2))
            seconds = int(parts.group(3))
            millis = int(parts.group(4))
            
            return hours * 3600 + minutes * 60 + seconds + millis / 1000
        
        return 0
    
    def _parse_json_transcript(self, content: str) -> List[Dict[str, Any]]:
        """Parse JSON transcript into segments"""
        data = json.loads(content)
        
        if "segments" in data:
            return data["segments"]
        elif "content" in data and isinstance(data["content"], list):
            return data["content"]
        else:
            # Try to extract from other structures
            segments = []
            # Implementation would depend on specific JSON structure
            return segments
    
    def _segments_to_transcript(self, segments: List[Dict[str, Any]]) -> str:
        """Convert segments back to transcript format"""
        lines = []
        
        for segment in segments:
            # Add timestamp
            time_str = self._seconds_to_time(segment["start"])
            line = f"[{time_str}] "
            
            # Add speaker if available
            if segment.get("speaker"):
                line += f"{segment['speaker']}: "
            
            # Add text
            line += segment["text"]
            
            lines.append(line)
        
        return '\n\n'.join(lines)